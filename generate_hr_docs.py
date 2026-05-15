"""Generate synthetic HR policy Word documents for SharePoint upload."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = r"C:\HRAgentService\sharepoint-docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def add_header(doc, title):
    doc.add_heading(title, level=0)
    p = doc.add_paragraph()
    p.add_run("Contoso Corporation | Human Resources").bold = True
    p.add_run("\nConfidential – Internal Use Only")
    p.add_run(f"\nEffective Date: January 1, 2026 | Last Revised: May 1, 2026")
    doc.add_paragraph()


# ─────────────────────────────────────────────
# Document 1: Personal Data Change – ESS Guide
# ─────────────────────────────────────────────
def create_ess_guide():
    doc = Document()
    set_style(doc)
    add_header(doc, "Personal Data Change – Employee Self-Service (ESS) Guide")

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "Contoso employees can update certain personal information directly through "
        "the Employee Self-Service (ESS) portal in Workday. These are considered "
        "'simple updates' and do not require HR approval."
    )

    doc.add_heading("Simple Updates (Self-Service via ESS)", level=1)
    doc.add_paragraph(
        "The following changes can be made directly by the employee in Workday ESS:"
    )
    items = [
        ("Emergency Contact", "Add, edit, or remove emergency contacts including name, relationship, phone, and email."),
        ("Home Contact Information", "Update home address, personal phone number, and personal email address."),
        ("Personal Information", "Update marital status, date of marriage, gender identity, pronouns, and ethnicity (voluntary)."),
        ("Preferred Name", "Change your preferred/display name (does not affect legal records). Visible in Teams, Outlook, and internal directory."),
    ]
    for title, desc in items:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        doc.add_paragraph("Steps:", style="List Bullet")
        doc.add_paragraph("1. Log in to Workday ESS Portal", style="List Number")
        doc.add_paragraph("2. Navigate to Personal Information > [Section]", style="List Number")
        doc.add_paragraph("3. Click 'Edit' and make your changes", style="List Number")
        doc.add_paragraph("4. Review and Submit", style="List Number")
        doc.add_paragraph("Changes take effect immediately. No manager approval required.")
        doc.add_paragraph()

    doc.add_heading("Life Event Examples", level=1)
    doc.add_paragraph(
        "If you've experienced a life event such as marriage, divorce, birth of a child, "
        "or domestic partnership, you may need to update multiple sections. Use the "
        "'Life Event' workflow in Workday to update everything in one flow."
    )
    doc.add_paragraph("Common life event scenarios:")
    scenarios = [
        "I got married → Update: Marital Status, Emergency Contact, Preferred Name (if applicable), Home Address (if moving)",
        "I moved to a new address → Update: Home Contact Information, Emergency Contact address",
        "I changed my phone number → Update: Home Contact, Emergency Contact",
        "I want to go by a nickname → Update: Preferred Name",
    ]
    for s in scenarios:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Processing Time", level=1)
    doc.add_paragraph("Simple ESS changes are processed instantly in the system.")

    doc.save(os.path.join(OUTPUT_DIR, "01_Personal_Data_Change_ESS_Guide.docx"))
    print("Created: 01_Personal_Data_Change_ESS_Guide.docx")


# ─────────────────────────────────────────────
# Document 2: Complex Personal Data Changes
# ─────────────────────────────────────────────
def create_complex_changes():
    doc = Document()
    set_style(doc)
    add_header(doc, "Complex Personal Data Changes – Policy & Procedures")

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "Certain personal data changes require verification documentation and HR review. "
        "These are classified as 'complex updates' and must be submitted through the "
        "HR Service Center with supporting evidence."
    )

    doc.add_heading("Complex Change Categories", level=1)

    changes = [
        {
            "title": "Legal Name Change",
            "desc": "Required when an employee's legal name changes due to marriage, divorce, court order, or gender transition.",
            "docs": "Government-issued ID showing new legal name (e.g., updated passport, court order, marriage certificate)",
            "process": "Submit request via HR Service Center > Attach documentation > HR reviews within 3-5 business days > Updated in payroll and all systems",
            "link": "https://contoso-workday.example.com/legal-name-change"
        },
        {
            "title": "Passport & Visa Updates",
            "desc": "Required when passport or visa information changes, especially for work authorization purposes.",
            "docs": "Copy of new passport/visa, I-9 re-verification form (if applicable)",
            "process": "Submit via HR Service Center > Immigration team reviews > Updated within 5-7 business days",
            "link": "https://contoso-workday.example.com/passport-visa-update"
        },
        {
            "title": "Government ID Changes",
            "desc": "Updates to Social Security Number, Tax ID, or National Insurance Number.",
            "docs": "Official government correspondence or new SSN card",
            "process": "Submit via HR Service Center > Payroll team verifies > Updated in next pay cycle",
            "link": "https://contoso-workday.example.com/govt-id-update"
        },
        {
            "title": "Licenses & Other IDs",
            "desc": "Professional licenses, certifications, or company badge updates.",
            "docs": "Copy of new license/certification",
            "process": "Submit via HR Service Center > HR verifies > Updated within 3 business days",
            "link": "https://contoso-workday.example.com/licenses-update"
        },
        {
            "title": "Payment Election (Bank Details)",
            "desc": "Changes to direct deposit, bank account, or payment method.",
            "docs": "Void check or bank letter confirming account details",
            "process": "Submit via HR Service Center > Payroll verifies > Effective next pay period (allow 1 full pay cycle)",
            "link": "https://contoso-workday.example.com/payment-election"
        },
        {
            "title": "Photo Change",
            "desc": "Update your official company photo (badge, directory, Teams).",
            "docs": "Professional headshot meeting company guidelines (min 400x400px, neutral background)",
            "process": "Submit via HR Service Center > Approved within 2 business days",
            "link": "https://contoso-workday.example.com/photo-update"
        },
    ]

    for c in changes:
        doc.add_heading(c["title"], level=2)
        doc.add_paragraph(c["desc"])
        doc.add_paragraph(f"Required Documentation: {c['docs']}")
        doc.add_paragraph(f"Process: {c['process']}")
        doc.add_paragraph(f"Workday Link: {c['link']}")
        doc.add_paragraph()

    doc.add_heading("Important Notes", level=1)
    doc.add_paragraph("• All complex changes require valid supporting documentation")
    doc.add_paragraph("• HR reserves the right to request additional verification")
    doc.add_paragraph("• Changes affecting payroll must be submitted by the 15th of the month to take effect in the current cycle")
    doc.add_paragraph("• For urgent requests, contact HR Service Center at hr-help@contoso.com")

    doc.save(os.path.join(OUTPUT_DIR, "02_Complex_Personal_Data_Changes_Policy.docx"))
    print("Created: 02_Complex_Personal_Data_Changes_Policy.docx")


# ─────────────────────────────────────────────
# Document 3: Grievance Policy & ERLR Procedures
# ─────────────────────────────────────────────
def create_grievance_policy():
    doc = Document()
    set_style(doc)
    add_header(doc, "Employee Grievance Policy & ERLR Intake Procedures")

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This policy establishes the framework for employees to raise workplace concerns, "
        "complaints, and grievances. It defines the distinction between general workplace "
        "concerns and formal grievances that require Employee Relations & Labor Relations (ERLR) intervention."
    )

    doc.add_heading("Scope", level=1)
    doc.add_paragraph("This policy applies to all Contoso employees worldwide, including full-time, part-time, and contingent workers.")

    doc.add_heading("Definitions", level=1)
    doc.add_heading("Formal Grievance (ERLR-Routed)", level=2)
    doc.add_paragraph(
        "A formal complaint involving serious workplace misconduct that requires investigation "
        "by the Employee Relations team. These matters are confidential and follow a structured intake process."
    )
    doc.add_paragraph("Examples of formal grievances:")
    serious = [
        "Workplace harassment (sexual, racial, disability-based, or other protected characteristics)",
        "Discrimination based on age, gender, race, religion, disability, sexual orientation, or national origin",
        "Retaliation against an employee for reporting concerns or participating in an investigation",
        "Bullying or intimidation that creates a hostile work environment",
        "Threats of violence or actual workplace violence",
        "Ethical violations or fraud",
        "Wage theft or deliberate withholding of compensation",
        "Unsafe working conditions that management has failed to address",
        "Violation of accommodation requests (ADA, religious, pregnancy)",
    ]
    for s in serious:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Grievance Out-of-Scope (GOOS)", level=2)
    doc.add_paragraph(
        "Workplace concerns that, while valid, do not meet the threshold for formal ERLR investigation. "
        "These should be resolved through normal management channels, team discussions, or mediation."
    )
    doc.add_paragraph("Examples of GOOS topics:")
    goos = [
        "A coworker moved my chair or desk items",
        "Disagreement about meeting room bookings",
        "Preference for different office temperature",
        "Colleague is too loud on calls",
        "Disagreement about project approach or priorities",
        "Parking spot disputes",
        "Kitchen/break room cleanliness issues",
        "Colleague's personal habits (e.g., eating at desk)",
        "Schedule preference conflicts (not related to accommodation)",
        "One-time rude comment that was not discriminatory in nature",
    ]
    for g in goos:
        doc.add_paragraph(g, style="List Bullet")

    doc.add_heading("How to Determine: Grievance vs. GOOS", level=1)
    doc.add_paragraph("Ask these screening questions:")
    doc.add_paragraph("1. Does the concern involve a protected characteristic (race, gender, age, disability, religion, etc.)?", style="List Number")
    doc.add_paragraph("2. Is there a pattern of behavior (not a one-time minor incident)?", style="List Number")
    doc.add_paragraph("3. Has the employee felt unsafe or threatened?", style="List Number")
    doc.add_paragraph("4. Could this constitute illegal activity?", style="List Number")
    doc.add_paragraph("5. Has the employee already tried to resolve it with their manager without success?", style="List Number")
    doc.add_paragraph()
    doc.add_paragraph(
        "If the answer to ANY of questions 1-4 is YES, route to ERLR formal intake. "
        "If only question 5 is YES, suggest mediation or skip-level conversation first."
    )

    doc.add_heading("Formal Grievance Intake Process", level=1)
    doc.add_paragraph("1. Employee submits initial concern to HR Concierge or HR Service Center")
    doc.add_paragraph("2. HR Concierge screens for GOOS vs. formal grievance")
    doc.add_paragraph("3. If formal grievance: Employee is directed to the ERLR Intake Form")
    doc.add_paragraph("4. ERLR Intake Form collects: description of incident(s), dates, witnesses, desired outcome")
    doc.add_paragraph("5. ERLR case manager assigned within 48 hours")
    doc.add_paragraph("6. Investigation conducted (typically 10-30 business days)")
    doc.add_paragraph("7. Resolution communicated to all parties")
    doc.add_paragraph()
    doc.add_paragraph("ERLR Intake Form Link: https://contoso-workday.example.com/erlr-intake")

    doc.add_heading("Confidentiality", level=1)
    doc.add_paragraph(
        "All grievance reports are treated with strict confidentiality. Information is shared "
        "only on a need-to-know basis. Retaliation against anyone who files a grievance is "
        "strictly prohibited and will result in disciplinary action."
    )

    doc.add_heading("Non-Retaliation", level=1)
    doc.add_paragraph(
        "Contoso prohibits retaliation against any employee who in good faith reports a concern, "
        "files a grievance, or participates in an investigation. If you believe you are experiencing "
        "retaliation, report it immediately through the ERLR intake process."
    )

    doc.save(os.path.join(OUTPUT_DIR, "03_Grievance_Policy_ERLR_Procedures.docx"))
    print("Created: 03_Grievance_Policy_ERLR_Procedures.docx")


# ─────────────────────────────────────────────
# Document 4: GOOS Resolution Guide
# ─────────────────────────────────────────────
def create_goos_guide():
    doc = Document()
    set_style(doc)
    add_header(doc, "Grievance Out-of-Scope (GOOS) – Resolution Guide for Managers")

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This guide helps managers and HR representatives handle workplace concerns that "
        "fall outside the scope of formal ERLR grievance processes. These 'GOOS' issues "
        "are best resolved through communication, mediation, and team norms."
    )

    doc.add_heading("Resolution Approaches", level=1)

    approaches = [
        ("Direct Conversation", "Encourage the employee to have a direct, respectful conversation with the other party. Provide coaching on how to approach it."),
        ("Manager Mediation", "The manager facilitates a discussion between parties to find a mutually acceptable solution."),
        ("Team Norms Agreement", "Establish clear team norms about shared spaces, noise levels, scheduling, etc."),
        ("Facilities Request", "For environmental issues (temperature, noise, seating), submit a facilities ticket."),
        ("Skip-Level Meeting", "If the concern involves the direct manager, the employee may request a skip-level discussion."),
    ]

    for title, desc in approaches:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_heading("Escalation Path", level=1)
    doc.add_paragraph(
        "If a GOOS issue persists after multiple resolution attempts, or if the employee "
        "provides new information suggesting the concern may actually meet formal grievance "
        "criteria, it can be escalated to ERLR for re-evaluation."
    )

    doc.add_heading("Documentation", level=1)
    doc.add_paragraph(
        "While GOOS issues don't require formal documentation, managers should keep brief "
        "notes of conversations and agreed-upon actions for reference."
    )

    doc.save(os.path.join(OUTPUT_DIR, "04_GOOS_Resolution_Guide.docx"))
    print("Created: 04_GOOS_Resolution_Guide.docx")


# ─────────────────────────────────────────────
# Document 5: HR Service Catalog
# ─────────────────────────────────────────────
def create_service_catalog():
    doc = Document()
    set_style(doc)
    add_header(doc, "HR Service Catalog – Employee Quick Reference")

    doc.add_heading("HR Services Overview", level=1)
    doc.add_paragraph(
        "This document lists all HR services available to Contoso employees, "
        "including self-service options and those requiring HR assistance."
    )

    doc.add_heading("Personal Data Changes", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Change Type"
    hdr[1].text = "Self-Service?"
    hdr[2].text = "Approval Required?"
    hdr[3].text = "Processing Time"

    data = [
        ("Emergency Contact", "Yes", "No", "Immediate"),
        ("Home Address", "Yes", "No", "Immediate"),
        ("Personal Info (marital status)", "Yes", "No", "Immediate"),
        ("Preferred Name", "Yes", "No", "Immediate"),
        ("Legal Name", "No - HR Assisted", "Yes + Docs", "3-5 days"),
        ("Passport/Visa", "No - HR Assisted", "Yes + Docs", "5-7 days"),
        ("Government ID (SSN/Tax ID)", "No - HR Assisted", "Yes + Docs", "Next pay cycle"),
        ("Bank Details", "No - HR Assisted", "Yes + Docs", "Next pay period"),
        ("Photo", "No - HR Assisted", "Yes", "2 days"),
        ("Licenses/Certifications", "No - HR Assisted", "Yes + Docs", "3 days"),
    ]
    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_heading("Grievance & Employee Relations", level=1)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Service"
    hdr2[1].text = "Channel"
    hdr2[2].text = "Response Time"

    data2 = [
        ("Formal Grievance Filing", "ERLR Intake Form", "Case assigned within 48 hrs"),
        ("Workplace Concern (non-grievance)", "Manager / HR Concierge", "Same day guidance"),
        ("Mediation Request", "HR Service Center", "Scheduled within 5 days"),
        ("Anonymous Reporting", "Ethics Hotline", "Acknowledged within 24 hrs"),
    ]
    for row_data in data2:
        row = table2.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_heading("Contact Information", level=1)
    doc.add_paragraph("HR Service Center: hr-help@contoso.com")
    doc.add_paragraph("ERLR Team: erlr@contoso.com")
    doc.add_paragraph("Ethics Hotline: 1-800-CONTOSO (anonymous)")
    doc.add_paragraph("HR Concierge (AI Assistant): Available 24/7 via Teams or Portal")

    doc.save(os.path.join(OUTPUT_DIR, "05_HR_Service_Catalog.docx"))
    print("Created: 05_HR_Service_Catalog.docx")


if __name__ == "__main__":
    create_ess_guide()
    create_complex_changes()
    create_grievance_policy()
    create_goos_guide()
    create_service_catalog()
    print(f"\nAll documents saved to: {OUTPUT_DIR}")
    print("Upload these to your SharePoint HR site document library.")
